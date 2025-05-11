import base64
import json
import os
import logging
from channels.layers import get_channel_layer
from asgiref.sync import async_to_sync
from keydev_reports.models import ReportTemplate
import qrcode
from io import BytesIO
from docx import Document as xDocument
from docx.shared import Inches
from python_docx_replace.paragraph import Paragraph
from django.conf import settings

from project.users.models import User
from project.documents.tasks import get_template_report
from project.documents.proxy_models.position_proxy import UserEmployeeProxy
from project.users.models import UserEmployee

from .models import Document, Signature

def create_document(
    name: str,
    created_user: User,
    responsible_users: list[User],
    description: str = None,
    file: bytes = None,
    document_template: ReportTemplate = None,
) -> Document:
    doc = Document.objects.create(
        name=name,
        created_by=created_user,
        description=description,
        file=file,
        document_template=document_template,
    )
    doc.responsible_users.set(responsible_users)
    if document_template:
        create_document_from_template(
            user=created_user,
            document_id=doc.id,
            template_id=document_template.pk
        )
    return doc


def document_user_sign(document_id: int, user_id: int, password: str):
    document = Document.objects.filter(id=document_id).first()
    if document is None:
        raise Exception("Document not found")
    user = User.objects.filter(id=user_id).first()
    if user is None:
        raise Exception("User not found")
    if document.file is None:
        raise Exception("Document file not found")
    with document.file.open('rb') as f:
        file_bytes = f.read()
    signature = user.sign_document(file_bytes, password)
    result = user.verify_signature(file_bytes, signature)
    if result:
        signature_b64 = base64.b64encode(signature).decode('utf-8')
        Signature.objects.create(
            document=document,
            user=user,
            signature=signature_b64,
        )
        if document.responsible_users.count() == Signature.objects.filter(document=document).count():
            try:
                set_qrcode_to_document(document)
            except Exception as e:
                logging.error(f'Failed to set qrcode to document with id {document.id}: {str(e)}')
            else:
                logging.info(f'Qrcode successfully set to document with id {document.id}')
        return True
    return False


def notify_users_about_document(user: User, document: Document):
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"user_{user.pk}",
        {
            'type': 'document_created',
            'document_id': document.pk,
            'title': document.name,
        }
    )


def create_document_from_template(
    user: User,
    document_id: int,
    template_id: int,
) -> dict:
    employee = UserEmployee.objects.filter(user=user).first()
    if employee is None:
        raise Exception('У пользователя нет должности!')
    proxy_obj = UserEmployeeProxy.proxy_objects.get_all_data().get(pk=employee.pk)
    report_data = proxy_obj.get_report_data()
    result = get_template_report.delay(
        report_data=report_data,
        user_name=user.username,
        report_id=template_id,
        provider_name='ProxyModelProvider',
        document_id=document_id
    )
    return {
        'task_id': result.id,
        'status': result.status
    }


def set_qrcode_to_document(
    document: Document
) -> None:
    signatures = Signature.objects.select_related('user').filter(document=document)
    sign_data = []
    if signatures.count() < 1:
        raise Exception('Для документа не была найдена ни одна подпись.')
    for signature in signatures:
        sign_data.append(
            {
                'username': signature.user.username,
                'signature': signature.signature
            }
        )
    json_data = json.dumps(sign_data)
    qr_img = qrcode.make(json_data)
    qr_buffer = BytesIO()
    qr_img.save(qr_buffer, format='PNG')
    qr_buffer.seek(0)
    file_path = os.path.join(settings.MEDIA_ROOT, document.file.name)
    doc = xDocument(file_path)
    find = False
    for para in Paragraph.get_all(doc):
        if '{{QR}}' in para.text:
            parent = para._element.getparent()
            parent.remove(para._element)
            new_para = doc.add_paragraph()
            run = new_para.add_run()
            run.add_picture(qr_buffer, width=Inches(0.787))
            find = True
            break
    if not find:
        raise Exception('В документе нет метки {{QR}}')
    doc.save(file_path)








